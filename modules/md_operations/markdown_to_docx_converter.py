"""
Markdown to DOCX converter module.
Converts markdown files to Word documents with proper style mapping.
"""

import re
from pathlib import Path
from html.parser import HTMLParser
from typing import Optional

from docx import Document


# Regex to match invalid XML characters (control chars except tab, newline, carriage return)
INVALID_XML_CHARS_RE = re.compile(
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]"
)


def _sanitize_text(text: str) -> str:
    """Remove invalid XML characters from text."""
    return INVALID_XML_CHARS_RE.sub("", text)


from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
from markdown.extensions.tables import TableExtension
from markdown.extensions.fenced_code import FencedCodeExtension


class DocxBuilder(HTMLParser):
    """Parse HTML from markdown and build DOCX document."""

    def __init__(self, document: Document):
        super().__init__()
        self.doc = document
        self.current_paragraph = None
        self.current_run = None
        self.style_stack: list[str] = []
        self.list_stack: list[str] = []  # 'ul' or 'ol'
        self.list_counters: list[int] = []
        self.in_code_block = False
        self.code_content = ""
        self.table_data: list[list[str]] = []
        self.current_row: list[str] = []
        self.current_cell = ""
        self.in_table = False
        self.in_blockquote = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, Optional[str]]]):
        tag = tag.lower()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            self.current_paragraph = self.doc.add_heading(level=level)
            self.style_stack.append(tag)

        elif tag == "p":
            if not self.in_table:
                self.current_paragraph = self.doc.add_paragraph()
                if self.in_blockquote:
                    self._set_paragraph_indent(self.current_paragraph, Pt(36))

        elif tag in ("strong", "b"):
            self.style_stack.append("bold")

        elif tag in ("em", "i"):
            self.style_stack.append("italic")

        elif tag == "code":
            if self.in_code_block:
                pass  # Already in code block, ignore nested code tag
            else:
                self.style_stack.append("code")

        elif tag == "pre":
            self.in_code_block = True
            self.code_content = ""

        elif tag == "ul":
            self.list_stack.append("ul")

        elif tag == "ol":
            self.list_stack.append("ol")
            self.list_counters.append(0)

        elif tag == "li":
            self.current_paragraph = self.doc.add_paragraph()
            if self.list_stack:
                list_type = self.list_stack[-1]
                indent_level = len(self.list_stack) - 1
                if list_type == "ul":
                    self.current_paragraph.style = "List Bullet"
                else:
                    self.current_paragraph.style = "List Number"
                    if self.list_counters:
                        self.list_counters[-1] += 1
                self._set_paragraph_indent(
                    self.current_paragraph, Pt(18 * indent_level)
                )

        elif tag == "blockquote":
            self.in_blockquote = True

        elif tag == "table":
            self.in_table = True
            self.table_data = []

        elif tag == "tr":
            self.current_row = []

        elif tag in ("td", "th"):
            self.current_cell = ""

        elif tag == "hr":
            self._add_horizontal_rule()

    def handle_endtag(self, tag: str):
        tag = tag.lower()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            if self.style_stack and self.style_stack[-1] == tag:
                self.style_stack.pop()
            self.current_paragraph = None

        elif tag == "p":
            self.current_paragraph = None

        elif tag in ("strong", "b"):
            if self.style_stack and self.style_stack[-1] == "bold":
                self.style_stack.pop()

        elif tag in ("em", "i"):
            if self.style_stack and self.style_stack[-1] == "italic":
                self.style_stack.pop()

        elif tag == "code":
            if not self.in_code_block and self.style_stack:
                if self.style_stack[-1] == "code":
                    self.style_stack.pop()

        elif tag == "pre":
            if self.in_code_block:
                self._add_code_block(self.code_content.strip())
                self.in_code_block = False
                self.code_content = ""

        elif tag == "ul":
            if self.list_stack and self.list_stack[-1] == "ul":
                self.list_stack.pop()

        elif tag == "ol":
            if self.list_stack and self.list_stack[-1] == "ol":
                self.list_stack.pop()
            if self.list_counters:
                self.list_counters.pop()

        elif tag == "li":
            self.current_paragraph = None

        elif tag == "blockquote":
            self.in_blockquote = False

        elif tag == "table":
            self._build_table()
            self.in_table = False
            self.table_data = []

        elif tag == "tr":
            if self.current_row:
                self.table_data.append(self.current_row)
            self.current_row = []

        elif tag in ("td", "th"):
            self.current_row.append(self.current_cell.strip())
            self.current_cell = ""

    def handle_data(self, data: str):
        # Sanitize text to remove invalid XML characters
        data = _sanitize_text(data)

        if self.in_code_block:
            self.code_content += data
            return

        if self.in_table:
            self.current_cell += data
            return

        if not self.current_paragraph:
            # Skip whitespace outside paragraphs
            if data.strip():
                self.current_paragraph = self.doc.add_paragraph()
            else:
                return

        # Apply formatting based on style stack
        run = self.current_paragraph.add_run(data)

        if "bold" in self.style_stack:
            run.bold = True
        if "italic" in self.style_stack:
            run.italic = True
        if "code" in self.style_stack:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    def _add_code_block(self, code: str):
        """Add a shaded code block paragraph."""
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

        # Set shading via XML
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F0F0F0")
        paragraph._p.get_or_add_pPr().append(shading)

    def _add_horizontal_rule(self):
        """Add a horizontal rule as a paragraph with bottom border."""
        paragraph = self.doc.add_paragraph()
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _set_paragraph_indent(self, paragraph, indent):
        """Set left indent for a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(int(indent)))
        pPr.append(ind)

    def _build_table(self):
        """Build a DOCX table from collected table data."""
        if not self.table_data:
            return

        rows = len(self.table_data)
        cols = max(len(row) for row in self.table_data) if self.table_data else 0

        if rows == 0 or cols == 0:
            return

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(self.table_data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < len(row.cells):
                    row.cells[j].text = cell_text


def convert_markdown_to_docx(input_path: Path, output_path: Path) -> tuple[bool, str]:
    """
    Convert markdown file to DOCX with style mapping.

    Args:
        input_path: Path to input markdown file
        output_path: Path for output DOCX file

    Returns:
        Tuple of (success: bool, message: str)
    """
    try:
        # Read markdown content
        try:
            content = input_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            content = input_path.read_text(encoding="latin-1")

        # Sanitize content to remove invalid XML characters
        content = _sanitize_text(content)

        # Remove image references (skip images as per requirements)
        content = re.sub(r"!\[.*?\]\(.*?\)", "", content)

        # Convert markdown to HTML
        md = markdown.Markdown(
            extensions=[
                TableExtension(),
                FencedCodeExtension(),
                "nl2br",
            ]
        )
        html = md.convert(content)

        # Create DOCX document
        document = Document()

        # Parse HTML and build DOCX
        builder = DocxBuilder(document)
        builder.feed(html)

        # Save document
        document.save(str(output_path))

        return True, f"Successfully converted to {output_path.name}"

    except Exception as e:
        return False, str(e)
